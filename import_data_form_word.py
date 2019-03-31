#coding:utf-8
# @Info: 从人事档案表中提取数据
# @Author:Netfj@sina.com @File:import_data_form_word.py @Time:2019/3/30 6:50

import docx,logging
def logset():
    filename='runinfo.log'
    style='simplify'
    level_grade = logging.DEBUG
    LOG_FORMAT = "%(asctime)s|%(filename)s(%(lineno)d)%(funcName)s" \
                   "[%(levelname)s]%(message)s"
    DATE_FORMAT = "%m.%d.%H:%M"
    logging.basicConfig(filename=filename,
                        level=level_grade,
                        format=LOG_FORMAT,
                        datefmt=DATE_FORMAT)
logset()

class pickup_emp():
    db_table_0_a = {
        'name':'姓名', 'gender':'性别', 'birthday':'出生年月',
        'nation':'民族', 'native':'籍贯','birthplace':'出生地',
        'party_time':'入党时间', 'work_time':'参加工作时间','health':'健康状况',
        'profession':'专业技术职务','speciality':'熟悉专业有何专长',
        'education1':'全日制教育', 'academy1':'毕业院校系及专业',
        'education2':'在职教育', 'academy2':'毕业院校系及专业',
        'post_now':'现任职务',
        'post_will':'拟任职务',
        'post_remove':'拟免职务',
        'end':'简历'
    }
    db_table_0_b = {
        'resume_time':'简历时间','resume_post':'简历岗位'
    }

    def __init__(self):
        pass

    # 内部数据初始化
    def initialize(self):
        self.docx_file = ''             # 导入的 WORD 文件名：sample.docx
        self.tables_info = {}           # WORD文档中表的信息：几个几行几列
        self.table_items = {}           # 提取的表的初始数据：原始信息
        self.table_items_clean = {}     # 清洗处理过后的数据：可用信息
        self.import_to_db_data = {}     # 将要导入到数据表中的数据：精准信息
        logging.debug('内部数据初始化')

    # 更换文件
    def updata_word_file(self,docx_file):
        self.initialize()                   # 将内部数据初始化
        self.docx_file = docx_file          # 更换文件名
        doc = docx.Document(self.docx_file) # 打开文件
        self.tables = doc.tables            # 获取文件中的表格集
        self.tables_info.update(word文件名=self.docx_file)
        self.tables_info.update(表格数量=len(self.tables))
        for n in range(0,len(self.tables)):
            row = len(self.tables[n].rows)
            col = len(self.tables[n].columns)
            self.tables_info.update({'表格{}的行列数量（行，列）'.format(n):(row,col)})
        logging.info(self.tables_info)
        
    # 读取单元格的文本，参数：表序号、行号、列号
    def get_cell_text(self,table=0,row=0,col=0):
        t = self.tables[table].cell(row, col).text
        return t

    # 提取表的数据，保存在：self.table_items
    def import_data_from_word_file(self):
        for n in range(0,len(self.tables)):
            row = len(self.tables[n].rows)
            col = len(self.tables[n].columns)
            self.table_items[n]=[]
            for r in range(0,row):
                for c in range(0,col):
                    t = self.get_cell_text(table=n,row=r,col=c)
                    self.table_items[n].append(t)
        logging.info(self.table_items)

    def clean_data(self):
        # 清洗数据：1.将（简历之前）连续的重复的数据去除；
        #          2.简历之后数据不作处理；
        #          3.空数据删除,删除头尾换行符'\n'
        # 数据源：self.table_items[0]
        # 生成后保存的位置：self.table_items_clean.update['db_table_0_a']

        data = [i for i in self.table_items[0] if i != '']
        data = [i.replace(' ','').replace('　','').strip('\n') for i in data]
        last = ''
        del_items = []
        for n in range(0,len(data)):
            now = data[n]
            # print(n,'|',last,'|',now)
            if now == last:
                del_items.append(n)
            last = now

        data_clean = []
        for i in range(0,len(data)):
            if not i in del_items:
                data_clean.append(data[i])

        # 保存清理过的数据：表1的第1部分（简历之前）
        self.table_items_clean.update(db_table_0_a = data_clean)

    # 生成可供导入数据表的精准数据，保存位置：import_to_db_data
    def create_db_data(self):
        data_source = self.table_items_clean['db_table_0_a']
        print(data_source)
        items_db_colume = self.db_table_0_a
        print(items_db_colume)

        xm = {}
        for k in items_db_colume:
            xm.update({k:[items_db_colume[k]]})
            # TODO
            # print(k,items_db_colume[k])
        print(xm)


    def run(self):
        self.import_data_from_word_file()   # 导入数据
        self.clean_data()                   # 清洗数据
        self.create_db_data()               # 生成可供导入数据表的精准数据


filelist = ['emp_sample.docx','emp_xxb.docx']
w = pickup_emp()
w.updata_word_file(filelist[0])
w.run()



import sys
sys.exit()
