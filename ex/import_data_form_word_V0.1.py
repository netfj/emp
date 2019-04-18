#coding:utf-8
# @Info: 从人事档案表中提取数据
# @Author:Netfj@sina.com @File:word2db.py @Time:2019/3/30 6:50

import docx,logging
def logset():
    filename='runinfo.log'
    style='simplify'
    level_grade = logging.DEBUG
    LOG_FORMAT = "%(asctime)s|%(filename)s(%(lineno)d)%(funcName)s" \
                   "[%(levelname)s]%(message)s"
    DATE_FORMAT = "%m.%d.%H:%M"
    logging.basicConfig(filename=filename,
                        level=level_grade,
                        format=LOG_FORMAT,
                        datefmt=DATE_FORMAT)
logset()

class pickup_emp():
    db_table_0_a = {
        'name':'姓名', 'gender':'性别', 'birthday':'出生年月',
        'nation':'民族', 'native':'籍贯','birthplace':'出生地',
        'party_time':'入党时间', 'work_time':'参加工作时间','health':'健康状况',
        'profession':'专业技术职务','speciality':'熟悉专业有何专长',
        'education1':'全日制教育', 'academy1':'毕业院校系及专业',
        'education2':'在职教育', 'academy2':'毕业院校系及专业',
        'post_now':'现任职务',
        'post_will':'拟任职务',
        'post_remove':'拟免职务',
        'end':'简历'
    }
    db_table_0_b = {
        'resume_time':'简历时间','resume_post':'简历岗位'
    }

    def __init__(self):
        pass

    # 内部数据初始化
    def initialize(self):
        self.docx_file = ''             # 导入的 WORD 文件名：sample.docx
        self.tables_info = {}           # WORD文档中表的信息：几个几行几列
        self.table_items = {}           # 提取的表的初始数据：原始信息
        self.table_items_clean = {}     # 清洗处理过后的数据：可用信息
        self.import_to_db_data = {}     # 将要导入到数据表中的数据：精准信息
        logging.debug('内部数据初始化')

    # 更换文件
    def updata_word_file(self,docx_file):
        self.initialize()                   # 将内部数据初始化
        self.docx_file = docx_file          # 更换文件名
        doc = docx.Document(self.docx_file) # 打开文件
        self.tables = doc.tables            # 获取文件中的表格集
        self.tables_info.update(word文件名=self.docx_file)
        self.tables_info.update(表格数量=len(self.tables))
        for n in range(0,len(self.tables)):
            row = len(self.tables[n].rows)
            col = len(self.tables[n].columns)
            self.tables_info.update({'表格{}的行列数量（行，列）'.format(n):(row,col)})
        logging.info(self.tables_info)
        
    # 读取单元格的文本，参数：表序号、行号、列号
    def get_cell_text(self,table=0,row=0,col=0):
        t = self.tables[table].cell(row, col).text
        return t

    # 提取表的数据，保存在：self.table_items
    def import_data_from_word_file(self):
        for n in range(0,len(self.tables)):
            row = len(self.tables[n].rows)
            col = len(self.tables[n].columns)
            self.table_items[n]=[]
            for r in range(0,row):
                for c in range(0,col):
                    t = self.get_cell_text(table=n,row=r,col=c)
                    self.table_items[n].append(t)
        logging.info(self.table_items)

    def clean_data(self):
        # 清洗数据：1.将（简历之前）连续的重复的数据去除；
        #          2.简历之后数据不作处理；
        #          3.空数据删除,删除头尾换行符'\n'
        #          4.编程需要，暂时去除中间的 '\n'
        # 数据源：self.table_items[0]
        # 生成后保存的位置：self.table_items_clean.update['db_table_0_a']

        data = [i for i in self.table_items[0] if i != '']
        data = [i.replace(' ','').replace('　','').strip('\n').replace('\n','')
                for i in data]
        last = ''
        del_items = []
        for n in range(0,len(data)):
            now = data[n]
            # print(n,'|',last,'|',now)
            if now == last:
                del_items.append(n)
            last = now

        data_clean = []
        for i in range(0,len(data)):
            if not i in del_items:
                data_clean.append(data[i])

        # 保存清理过的数据：表1的第1部分（简历之前）
        self.table_items_clean.update(db_table_0_a = data_clean)
        logging.info(self.table_items_clean['db_table_0_a'])

    # 生成可供导入数据表的精准数据:处理表1上半部分（姓名到简历以前）
    # 数据来源：self.table_items_clean
    # 保存位置：self.import_to_db_data
    def create_db_data_t1a(self):
        # 生成定位列表
        items_db_colume = self.db_table_0_a
        xm = {}
        k = [ i for i in items_db_colume.keys()]
        v = [i for i in items_db_colume.values()]
        for n in range(0,len(k)-1):
            xm.update({k[n]:[v[n],v[n+1]]})

        '''
        print(xm)
        {'name': ['姓名', '性别'], 'gender': ['性别', '出生年月'], 
        'birthday': ['出生年月', '民族'], 'nation': ['民族', '籍贯'], 
        'native': ['籍贯', '出生地'], 'birthplace': ['出生地', '入党时间'],
         'party_time': ['入党时间', '参加工作时间'], 
         'work_time': ['参加工作时间', '健康状况'], 
         'health': ['健康状况', '专业技术职务'], 
         'profession': ['专业技术职务', '熟悉专业有何专长'],
          'speciality': ['熟悉专业有何专长', '全日制教育'], 
          'education1': ['全日制教育', '毕业院校系及专业'], 
          'academy1': ['毕业院校系及专业', '在职教育'], 
          'education2': ['在职教育', '毕业院校系及专业'], 
          'academy2': ['毕业院校系及专业', '现任职务'], 
          'post_now': ['现任职务', '拟任职务'], 
          'post_will': ['拟任职务', '拟免职务'], 
          'post_remove': ['拟免职务', '简历']}
          
        下面要做的工作，在清洗过后的数据 
            self.table_items_clean.update['db_table_0_a']
                (其格式为：
                    ['姓名', '张某某', '性别', '男', 
                    '出生年月\n(岁)', '1989.07\n（27岁）',
                     '民族', '汉族', '籍贯', '浙江宁波', ....])
        中，定位读取精准数据，如：
            name 是在 ‘姓名’ 与 ‘性别’之间的数据
        
        '''

        # 读取项目值
        data_result = []
        data_source = self.table_items_clean['db_table_0_a']
        for x in xm:
            # 读取项目。【参数】x：项目，xm[x]：项目的首尾区间
            t = self.create_db_data_t1a_readxm(data_source,x,xm[x])

            # 以元组形式保存项目信息，格式如：
            # [('name', '姓名', '张某某'), ('gender', '性别', '男'), .....]
            data_result.append((x,xm[x][0],t))

        # 保存数据
        self.import_to_db_data.update({'tb1a':data_result})
        print(self.import_to_db_data)



    def create_db_data_t1a_readxm(self,data_source,x,section):
        '''
        功能：被 create_db_data_t1a 调用，读取项目的值
        :param data_source: word文档中表的值（经过清洗)
        :param x:           项目（如 ‘姓名’ ）
        :param section:     项目的值所在的区间（如在 姓名 和 性别 之间）
        :return:            该项目的值（如：张某某）
        '''

        # 定位 section 0,其原理是：
        #  set(data_source[n]) & set(x) == x
        #  集合{data_source[n]} 与 集合 {x} 的值相符
        # 如： {'出生年月\n(岁)'} & {出生年月} == {'出','生','年','月'}
        for n in range(0,len(data_source)):
            if set(data_source[n]) & set(section[0]) == set(section[0]):
                list_number1 = n
                break

        # 定位 section 1
        for n in range(0,len(data_source)):
            if set(data_source[n]) & set(section[1]) == set(section[1]):
                list_number2 = n
                break

        if list_number2 - list_number1 > 2:
            logging.warning(
                '取值可能有问题：{}|{}|{}|区间{}-{}'
                    .format(self.docx_file, x, section[0],
                            list_number1, list_number2))

        t = ''
        for i in range(list_number1+1,list_number2):
            t += data_source[i]
            if list_number2 - list_number1 > 2:
                logging.debug('  -->{}|{}==> t:{}'.format(i,data_source[i],t))

        if list_number2 - list_number1 > 2:
            logging.debug('  ==>读取到的值(t):{}'.format(t))
        else:
            logging.debug('~~>读取到的值(t):{}'.format(t))

        return t


    def run(self):
        self.import_data_from_word_file()   # 导入数据
        self.clean_data()                   # 清洗数据

        # 生成可供导入数据表的精准数据
        self.create_db_data_t1a()      # 表1 上半部分（简历以前）         


filelist = ['emp_sample.docx','emp_xxb.docx']
w = pickup_emp()
w.updata_word_file(filelist[0])
w.run()

a = w.table_items_clean['db_table_0_a']
for i in range(0,len(a)):
    print(i,a[i])


import sys
sys.exit()
