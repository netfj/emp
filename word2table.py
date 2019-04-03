#coding:utf-8
# @Info: 从人事档案表中提取数据
# @Author:Netfj@sina.com @File:word2table.py @Time:2019/3/30 6:50

from flask_sqlalchemy import SQLAlchemy
from datetime import datetime
from setup_database import app, Person, Record_info
import docx,logging

def logset():
    filename='runinfo.log'
    level_grade = logging.DEBUG
    LOG_FORMAT = "%(asctime)s|%(filename)s(%(lineno)d)%(funcName)s[%(levelname)s]%(message)s"
    DATE_FORMAT = "%m.%d.%H:%M"
    logging.basicConfig(filename=filename,
                        level=level_grade,
                        format=LOG_FORMAT,
                        datefmt=DATE_FORMAT)
logset()

class pickup_emp():
    filelist = ['emp_sample.docx']
    db_table_0a = {
        'name':'姓名', 'gender':'性别', 'birthday':'出生年月',
        'nation':'民族', 'native':'籍贯','birthplace':'出生地',
        'party_time':'入党时间', 'work_time':'参加工作时间','health':'健康状况',
        'profession':'专业技术职务','speciality':'熟悉专业有何专长',
        'education1':'全日制教育', 'academy1':'毕业院校系及专业',
        'education2':'在职教育', 'academy2':'毕业院校系及专业',
        'post_now':'现任职务',
        'post_will':'拟任职务',
        'post_remove':'拟免职务'
    }
    db_table_0b = {
        'resume_time':'简历时间','resume_post':'简历岗位'
    }

    def __init__(self,filelist=None):
        if filelist:
            self.filelist = filelist
            self.run()

    # 内部数据初始化
    def initialize(self):
        self.docx_file = ''             # 导入的 WORD 文件名：sample.docx
        self.table_info = {}            # 提取的表的初始数据：原始信息
        self.table_info_clean = {}      # 清洗处理过后的数据：可用信息
        self.data2db = {}               # 将要导入到数据表中的数据：精准信息
        logging.debug('内部数据初始化')

    # 提取表的数据，保存在：self.table_items
    def import_data_from_word_file(self):
        items = []
        for n in range(0, len(self.tables)):
            row = len(self.tables[n].rows)
            col = len(self.tables[n].columns)
            for r in range(0, row):
                for c in range(0, col):
                    t = self.tables[n].cell(row, col).text

    # 更换文件
    def updata_word_file(self, docx_file):
        self.initialize()                       # 将内部数据初始化
        self.docx_file = docx_file              # 更换文件名
        doc = docx.Document(self.docx_file)     # 打开文件
        self.tables = doc.tables                # 获取文件中的表格集 句柄

        # 收集有关信息, 保存在：self.table_info['table_info']
        tb_info = {'文件名':self.docx_file}                     # 文件名
        tb_info.update({'表格数量':str(len(self.tables))})      # 表格数量
        for n in range(0,len(self.tables)):
            row = len(self.tables[n].rows)
            col = len(self.tables[n].columns)
            tb_info.update({'表格{}的行列数量（行，列）'.format(n):(row,col)})
        self.table_info.update(table_info = tb_info)

    # 提取表的原始数据，以 列表 形式，保存在-->字典：self.table_info['items_text']
    ''' 数据结构：列表，每项的含义{表的序号: [(行号，列号，内容TEXT), ...] 
    实例：{0: [(0, 0, '姓\u3000名'), (0, 1, '姓\u3000名'), 
            (0, 2, '张某某'), (0, 3, '性\u3000别'), (0, 4, '性\u3000别'), ...]
    '''
    def extract_from_word(self):

        d = {}
        for n in range(0, len(self.tables)):
            x = []
            row = len(self.tables[n].rows)
            col = len(self.tables[n].columns)
            for r in range(0, row):
                for c in range(0, col):
                    t = self.tables[n].cell(r, c).text
                    x.append((r,c,t))
            d.update({n:x})
        self.table_info.update({'items_text':d})
        logging.debug(self.table_info)  # 写入日志

    # 清洗数据  ： 1.有效取值，2.去空格、头尾换行符
    # 保存至字典： self.table_info_clean
    def extract_from_word_clean(self):
        # word文档和表的有关信息
        self.table_info_clean.update({'table_info':self.table_info['table_info']})

        # 处理表1的上半部分（简历之前）
        data = []
        data += self.get_line_0a(key='姓名')
        data += self.get_line_0a(key='民族')
        data += self.get_line_0a(key='入党时间')
        data += self.get_line_0a(key='专业技术职务')
        data += self.get_line_0a(key='全日制教育')
        data += self.get_line_0a(key='在职教育')
        data += self.get_line_0a(key='现任职务')
        data += self.get_line_0a(key='拟任职务')
        data += self.get_line_0a(key='拟免职务')
        # 保存清洗后的数据: 表1上半部分（简历之前）
        self.table_info_clean.update({'db_table_0a':data})
        logging.debug(self.table_info_clean)

    # 生成可以插入到数据库表中的数据：精准数据
    # 保存位置
    def data2db_create(self):
        # word文档和表的有关信息
        self.data2db.update({'table_info':self.table_info['table_info']})

        # 简历之前
        data = self.table_info_clean['db_table_0a']     # 读取清洗过后的表0上半部分数据
        d2b = {}    # 在这里存放将要存入数据库表中的精准数据
        xm = self.db_table_0a   # 表0上半部分的项目列表：姓名、性别、民族...
        for k in xm.keys():
            for index, value in enumerate(data):
                if set(value) & set(xm[k]) == set(xm[k]):
                    xm_value = data[index+1]
                    break
            d2b.update({k:xm_value})
        self.data2db.update({'db_table_0a':d2b})
        logging.debug(self.data2db)

    def get_line_0a(self, key='项目名称'):
        ds = self.table_info['items_text'][0]  # 提取的原始数据
        for cell in ds:
            if set(cell[2]) & set(key) == set(key):
                row = cell[0]
                break
        lt = []
        for cell in ds:
            if cell[0] == row:
                lt.append(cell[2].replace(' ', '').replace('　', ''))

        #去重
        lt2 = []
        for x in lt:
            if len(lt2) == 0:       # 新列表为空时，先放一个元素
                lt2.append(x)
                continue
            if x != lt2[-1]:        # 新追加时，查看与新列表最后一个元素是否相等
                lt2.append(x)

        return lt2

    def import_data_to_table(self):
        db = SQLAlchemy(app)

        # 将表的信息写入数据库
        try:
            # 表0 上半部分
            result = db.session.execute(Person.__table__.insert(),self.data2db['db_table_0a'])
            db.session.commit()


            # 将工作过程写入辅助表 record_info
            record_info = Record_info()
            record_info.id_person = result.lastrowid    # 增加记录的id号
            record_info.mode = 'word'
            record_info.info = self.docx_file
            record_info.data_souce = '{}'.format(self.table_info)
            record_info.data_clean = '{}'.format(self.table_info_clean)
            record_info.data2db = '{}'.format(self.data2db)
            record_info.dt = datetime.today()

            try:
                db.session.add(record_info)
                db.session.commit()
            except Exception as e:
                logging.error('写入工作过程错误')
                db.session.rollback()


        except Exception as e:
            logging.error('写入数据库错误:{}'.format(self.data2db['db_table_0a']))
            db.session.rollback()



    def run(self):
        logging.info('导入文件列表:'+  ('|').join(self.filelist))
        fn = len(self.filelist)
        for index,f in enumerate(self.filelist):
            msg = '进程: {}/{} — {}'.format(index+1,fn,f)
            logging.info(msg)
            print(msg)
            self.updata_word_file(f)        # 更换当前导入文件
            self.extract_from_word()        # 提取原始数据
            self.extract_from_word_clean()  # 清洗数据
            self.data2db_create()           # 提取精准数据（可以插入表中）
            self.import_data_to_table()     # 导入数据

            logging.info(msg + ' — 完成')

if __name__ == '__main__':
    filelist = ['emp_sample.docx','emp_xxb.docx']
    w = pickup_emp(filelist)



    import sys
    sys.exit()
