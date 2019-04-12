#coding:utf-8
# @Info: 从人事档案表中提取数据
# @Author:Netfj@sina.com @File:word2db.py @Time:2019/3/30 6:50

import docx,logging,os,re
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime
from setup_database import app, Person, Record_info, Home, Dwdm
from win32com import client
from tempfile import gettempdir
from random import randint
from shutil import rmtree
from time import sleep

def logset():
    filename='runinfo.log'
    level_grade = logging.DEBUG
    LOG_FORMAT = "%(asctime)s|%(filename)s(%(lineno)d)%(funcName)s[%(levelname)s]%(message)s"
    DATE_FORMAT = "%m.%d.%H:%M"
    logging.basicConfig(filename=filename,
                        level=level_grade,
                        format=LOG_FORMAT,
                        datefmt=DATE_FORMAT)
logset()

class pickup_emp():
    filelist = ['emp_sample.docx']
    db_table_0a = {
        'name':'姓名', 'gender':'性别', 'birthday':'出生年月',
        'nation':'民族', 'native':'籍贯','birthplace':'出生地',
        'party_time':'入党时间', 'work_time':'参加工作时间','health':'健康状况',
        'profession':'专业技术职务','speciality':'熟悉专业有何专长',
        'education1':'全日制教育', 'academy1':'毕业院校系及专业',
        'education2':'在职教育', 'academy2':'毕业院校系及专业',
        'post_now':'现任职务',
        'post_will':'拟任职务',
        'post_remove':'拟免职务'
    }
    db_table_0b = {
        'resume_time':'简历时间','resume_post':'简历岗位'
    }
    db_table_1a = {
        'reward':'奖惩情况','evaluation':'年度考核结果','reason':'任免理由'
    }
    db_table_1b = {
        'title':'称谓','name':'姓名','birthday':'出生年月',
        'party':'政治面貌','work':'工作单位及职务'
    }


    def __init__(self,filelist=None):
        if filelist:
            self.filelist = filelist        # 导入的文件列表
            self.set_tmp_path()             # 设置临时目录
            self.dwdm = []                  # 单位代码
            self.run_info = {'fault':[],'sucess':[]}     # 记录处理文件成功、失败

            self.db = SQLAlchemy(app)       # 新建一个实例
            self.word = client.Dispatch("Word.Application") # 启动word进程

            self.run()                      # 开始执行控制中心程序

            self.word.Quit()                # 退出 word
            self.db.close_all_sessions      # 关闭数据库连接
            self.clean_tmp_path()           # 清理临时目录

    def set_tmp_path(self):
        # 设置本类的临时目录：操作系统的临时目录 + 本系统特定的目录
        self.sys_tmp = os.path.join(gettempdir(), '_doc2docx')
        if not os.path.exists(self.sys_tmp):  # 临时目录：如果没有，则创建
            os.makedirs(self.sys_tmp)

    # 内部数据初始化
    def initialize(self):
        self.docx_file = ''             # 导入的 WORD 文件名：sample.docx
        self.imgData = ''               # 保存图片的二进制数据
        self.table_info = {}            # 提取的表的初始数据：原始信息
        self.table_info_clean = {}      # 清洗处理过后的数据：可用信息
        self.data2db = {}               # 将要导入到数据表中的数据：精准信息
        # logging.debug('内部数据初始化')

    def run_info_register(self,sucess=True,write_to_logfile = False):
        if write_to_logfile:

            total_file  = len(self.filelist)
            sucess_file = len(self.run_info['sucess'])
            fault_file  = len(self.run_info['fault'])

            msg0 = '【本次运行情况】总数{}：成功{}，失败{}'.format(total_file,sucess_file, fault_file)
            if fault_file>0:
                msg1 = msg0 + '{}。'.format(self.run_info['fault'])
            else:
                msg1 = msg0

            msg2 = msg0 + '。具体：{}'.format(self.run_info)

            print(msg1)
            logging.info(msg2)
            return True

        if sucess:
            self.run_info['sucess'].append(self.docx_file)
        else:
            self.run_info['fault'].append(self.docx_file)


    # doc --> docx
    def get_docx_file(self,word_file=''):
        if not word_file:
            msg = '没有传递文件！'
            print(msg)
            logging.error(msg)
            return False

        if not os.path.exists(os.path.realpath(word_file)):
            msg = '文件不存在: ' + word_file
            print(msg)
            logging.error(msg)
            return False

        if os.path.splitext(word_file)[1] == '.docx':
            return word_file

        else:
            tmp_docx_name = os.path.join(self.sys_tmp,
                          os.path.basename(word_file)+'.'+str(randint(0, 99999))+'.docx')

            open_word_file = os.path.realpath(word_file)

            try:  # doc ==> docx
                # word = client.Dispatch("Word.Application")
                doc = self.word.Documents.Open(open_word_file)
                doc.SaveAs(tmp_docx_name, 16)   #保存为docx
                doc.SaveAs(tmp_docx_name+'.html', 10)   #保存为html
                doc.Close()
                # word.Quit()
            except Exception as e:
                msg = '转换失败：{}'.format(e)
                print(msg)
                logging.error(msg)
                return False
            else:
                msg = '转换成功: {}'.format(tmp_docx_name)
                # print(msg)

                # 将图片写入 self.imgData 以备使用
                dir = tmp_docx_name+'.files'
                if os.path.isdir(dir):
                    list = os.listdir(dir)
                    if len(list)>0:
                        img_file = os.path.join(dir,list[0])
                        if os.path.isfile(img_file):
                            fp = open(img_file,'rb')
                            self.imgData = fp.read()
                            fp.close()

                return tmp_docx_name



    # 更换文件
    def updata_word_file(self, docx_file):
        self.initialize()                       # 将内部数据初始化
        self.docx_file = docx_file              # 更换文件名
        # logging.debug('开始处理：{}'.format(self.docx_file))

        open_word_file = self.get_docx_file(self.docx_file)
        if not open_word_file:      # 转换文件失败，则退出
            self.run_info_register(sucess=False)
            return False

        try:
            doc = docx.Document(open_word_file)     # 打开文件
        except Exception as e:
            msg = '打开文件({})错误：{}'.format(open_word_file,e)
            print(msg)
            logging.error(msg)
            return False
        else:
            msg = '打开文件成功!'

        self.tables = doc.tables                # 获取文件中的表格集 句柄
        if len(self.tables) !=2:
            msg = '{} 表格数量：{} (规定数量=2）'.format(self.docx_file,len(self.tables))
            print(msg)
            logging.error(msg)
            return False

        # 收集有关信息, 保存在：self.table_info['table_info']
        tb_info = {'文件名':self.docx_file}                     # 文件名
        tb_info.update({'表格数量':str(len(self.tables))})      # 表格数量
        for n in range(0,len(self.tables)):
            row = len(self.tables[n].rows)
            col = len(self.tables[n].columns)
            tb_info.update({'表格{}的行列数量（行，列）'.format(n):(row,col)})
        self.table_info.update(table_info = tb_info)

        # 读取图片
        for shape in doc.inline_shapes:
            contentID = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
            contentType = doc.part.related_parts[contentID].content_type
            if not contentType.startswith('image'):
                continue
            self.imgData = doc.part.related_parts[contentID]._blob
            break


        return True     # 成功


    # 更新单位页码库（部门、科室中队）代码
    def update_dwdm(self):
        file = re.sub('（[0-9]*人）','',self.docx_file)    #去掉（99人）的字样
        path2 = os.path.dirname(file)
        path1 = os.path.dirname(path2)
        p2 = path1[path1.rfind('\\') + 1:]
        p1 = path2[path2.rfind('\\') + 1:]
        dm_bm = p1[0:2]
        dm_ks = dm_bm + p2[0:2]
        mc_bm = p1[p1.find('.') + 1:]
        mc_ks = p2[p2.find('.') + 1:]
        self.dwdm.append((dm_bm,mc_bm))
        self.dwdm.append((dm_ks,mc_ks))


    # 将单位代码库 self.dwdm 写数数据库 employee.dwdms
    def update_db_dwdms(self):
        # 提取单位代码库中已经有的项目，防止重复冲突
        dwdms = Dwdm.query.all()
        dwdms_exist = [i.dm for i in dwdms]

        # 本次运行生成的，将要输入的库，先进行比较
        dwdm_new = [i
                    for i in self.dwdm
                    if i[0] not in dwdms_exist]

        # 将需要补充的 dwdm_new  写入数据库
        if len(dwdm_new)==0: return True
        logging.info('新增单位代码库：{}'.format(dwdm_new))
        try:
            # self.db.session.execute(Dwdm.__table__.insert(),dwdm_new)
            add_values = '{}'.format(dwdm_new)[1:-1]    #列表转换为字符，并且去掉首尾的 []
            sql = "insert into dwdms(dm,mc) values {}".format(add_values)
            self.db.session.execute(sql)
            self.db.session.commit()
        except Exception as e:
            msg = '写入单位代码失败：{}'.format(e)
            print(msg)
            logging.error(msg)
        else:
            pass

    # --------------------------------------------------------------------------
    # 提取表的原始数据，以 列表 形式，保存在-->字典：self.table_info['items_text']
    ''' 数据结构：列表，每项的含义{表的序号: [(行号，列号，内容TEXT), ...] 
    实例：{0: [(0, 0, '姓\u3000名'), (0, 1, '姓\u3000名'), 
            (0, 2, '张某某'), (0, 3, '性\u3000别'), (0, 4, '性\u3000别'), ...]
    '''
    def extract_from_word(self):
        d = {}
        for n in range(0, len(self.tables)):
            x = []
            row = len(self.tables[n].rows)
            col = len(self.tables[n].columns)
            for r in range(0, row):
                for c in range(0, col):
                    t = self.tables[n].cell(r, c).text
                    x.append((r,c,t))
            d.update({n:x})
        self.table_info.update({'items_text':d})
        logging.debug('原始信息：{}'.format(self.table_info))  # 写入日志

    # 清洗数据  ： 1.有效取值，2.视情况，去空格、头尾换行符
    # 保存至字典： self.table_info_clean
    def extract_from_word_clean(self):
        # word文档和表的有关信息
        self.table_info_clean.update({'table_info':self.table_info['table_info']})

        # 处理表0的上半部分（简历之前） ===============================
        data_0a = []
        data_0a += self.get_line(key='姓名')
        data_0a += self.get_line(key='民族')
        data_0a += self.get_line(key='入党时间')
        data_0a += self.get_line(key='专业技术职务')
        data_0a += self.get_line(key='全日制教育')
        data_0a += self.get_line(key='在职教育')
        data_0a += self.get_line(key='现任职务')
        data_0a += self.get_line(key='拟任职务')
        data_0a += self.get_line(key='拟免职务')

        # 修整：去空格
        data_0a_new = [ i.replace(' ', '').replace('　', '') for i in data_0a]

        # 保存清洗后的数据: 表0上半部分（简历之前）
        self.table_info_clean.update({'db_table_0a':data_0a_new})

        # 处理表0的下半部分（简历） =====================================
        data_0b = []
        data_0b += self.get_line(key='简历')

        # 修整：去首尾换行符 '\n'
        data_0b_new = [i.strip('\n') for i in data_0b]

        # 保存清洗后的数据: 表0下半部分（简历之前）
        self.table_info_clean.update({'db_table_0b': data_0b_new})

        # 处理表1的上半部分（奖惩、考核、任免理由） =======================
        data_1a = []
        data_1a += self.get_line(table_number = 1, key='奖惩情况')
        data_1a += self.get_line(table_number = 1, key='年度考核结果')
        data_1a += self.get_line(table_number = 1, key='任免理由')

        # 修整：去首尾换行符 '\n'
        data_1a_new = [i.strip('\n') for i in data_1a]

        # 保存清洗后的数据: 表1上半部分
        self.table_info_clean.update({'db_table_1a': data_1a_new})

        # 处理表1的 下 半部分（家庭主要成员） =======================
        data_1b =[]
        data_home = self.get_lines(key='家庭主要成员及重要社会关系')
        for key,value in data_home.items():
            if len(value)>=6:
                if not set(value[1]) & set('称谓') == set('称谓'):
                    data_1b.append({'title':value[1],'name':value[2],
                                   'birthday':value[3],'party':value[4],
                                    'work':value[5]})
        # 保存清洗后的数据: 表1下半部分
        self.table_info_clean.update({'db_table_1b': data_1b})

        # 将清洗后的数据，保存到日志 =================================
        logging.debug('有用信息：{}'.format(self.table_info_clean))


    def get_line(self, table_number = 0, key='项目名称'):
        '''
        从原始数据提取【不重复】的行，并作适当处理（去重）
        :param key: 起头的项目名称
        :return:    列表：起头的项目名称，该行的一行的数据
        '''
        ds = self.table_info['items_text'][table_number]  # 提取的原始数据
        for cell in ds:
            if set(cell[2]) & set(key) == set(key):
                row = cell[0]
                break
        lt = []
        for cell in ds:
            if cell[0] == row:
                lt.append(cell[2])

        # 去重
        lt2 = []
        for x in lt:
            if len(lt2) == 0:  # 新列表为空时，先放一个元素
                lt2.append(x)
                continue
            if x != lt2[-1]:  # 新追加时，查看与新列表最后一个元素是否相等
                lt2.append(x)

        return lt2

    def get_lines(self,key=None, table_number = 1):
        '''
        提取首字段是 key 的【所有行】。设计目的：用于提取表1【家庭主要成员及重要社会关系】
        :param key: 起头的字段名称，如 “家庭主要成员及重要社会关系”
        :return: 字典(各行的集合)，如：{列表1,列表2,...}
        '''
        result = {}
        ds = self.table_info['items_text'][table_number]  # 提取的原始数据
        for cell in ds:
            if set(cell[2]) & set(key) == set(key):
                row = cell[0]
                lt = []
                for cell in ds:
                    if cell[0] == row:
                        t = cell[2].replace('\n','')
                        if t!='':lt.append(t)
                result.update({row:lt})
        # 去重
        tmp = {}
        for key,value in result.items():
            lt = []
            [lt.append(i) for i in value if  i not in lt]
            if  lt not in tmp.values():
                tmp.update({key:lt})
        return tmp

    # 生成可以插入到数据库表中的数据：精准数据
    # 保存位置
    def data2db_create(self):
        # word文档和表的有关信息
        self.data2db.update({'table_info':self.table_info['table_info']})

        # 处理表0上半部分（简历之前）=============================
        # 读取清洗过后的表0上半部分数据
        data = self.table_info_clean['db_table_0a']
        d2b = {}    # 在这里存放将要存入数据库表中的精准数据
        xm = self.db_table_0a   # 表0上半部分的项目列表：姓名、性别、民族...
        for k in xm.keys():
            for index, value in enumerate(data):
                if set(value) & set(xm[k]) == set(xm[k]):
                    xm_value = data[index+1]
                    break
            d2b.update({k:xm_value})
        self.data2db.update({'db_table_0a':d2b})

        # 处理表0下半部分（简历）=============================
        # 读取清洗过后的表0下半部分数据（简历）:
        #   后面加一个空项（因为有的没有两列填写）
        data = self.table_info_clean['db_table_0b'] + ['']
        d2b = {
            'resume_time': data[1], 'resume_post': data[2]
        }
        self.data2db.update({'db_table_0b': d2b})

        # 处理表1上半部分（奖惩、考核、任免理由）===============
        data = self.table_info_clean['db_table_1a']
        d2b = {
            'reward': data[1], 'evaluation': data[3], 'reason': data[5]
        }
        self.data2db.update({'db_table_1a': d2b})

        # 处理表1下半部分（家庭主要成员）===============
            # 直接引用清洗后的数据
        d2b = self.table_info_clean['db_table_1b']
        self.data2db.update({'db_table_1b':d2b})


        # 写入日志 ==================================
        logging.debug('精确信息：{}'.format(self.data2db))

    def import_data_to_table(self):
        # db = SQLAlchemy(app)

        # 将表的信息写入数据库 ~~~~~~~~~~~~~~~~~~
        try:
            # 表0 上半部分（简历以前）
            result = self.db.session.execute(Person.__table__.insert(),
                                        self.data2db['db_table_0a'])

            # 表0 下半部分（简历）
            d2b = self.data2db['db_table_0b']
            self.db.session.query(Person).filter(Person.id == result.lastrowid ).update(d2b)

            # 表1 上半部分（奖惩、考核、任免）
            d2b = self.data2db['db_table_1a']
            self.db.session.query(Person).filter(Person.id == result.lastrowid ).update(d2b)

            # 写入图片
            if self.imgData != '':
                d2b = {'photo':self.imgData}
                self.db.session.query(Person).filter(Person.id == result.lastrowid ).update(d2b)


            # 表1 下半部分（家庭主要成员）
            d2b = self.data2db['db_table_1b']
            for n in range(0,len(d2b)):
                d2b[n].update({'id_person':result.lastrowid})
            self.db.session.execute(Home.__table__.insert(),d2b)

            # 写入
            self.db.session.commit()

            logging.debug('写入数据库：id:{}|word:{}'.format(result.lastrowid,self.docx_file))

            # 将工作过程写入辅助表 record_infos ：开始~~~~~~~~~~~~~~~~~~~~~
            record_info = Record_info()
            record_info.id_person = result.lastrowid  # 增加记录的id号
            record_info.mode = 'word'
            record_info.info = self.docx_file
            record_info.data_souce = '{}'.format(self.table_info)
            record_info.data_clean = '{}'.format(self.table_info_clean)
            record_info.data2db = '{}'.format(self.data2db)
            record_info.dt = datetime.today()
            try:
                self.db.session.add(record_info)
                self.db.session.commit()
            except Exception as e:
                msg = '写入工作过程错误:{} (数据表写入正常)'.format(e)
                logging.error(msg)
                print(msg)
                self.db.session.rollback()
            # 将工作过程写入辅助表 record_infos ：结束~~~~~~~~~~~~~~~~~~~~~

            return True


        except Exception as e:
            msg = '写入数据库错误:{}'.format(e)
            logging.error(msg)
            print(msg)
            self.db.session.rollback()

            return False



    def clean_tmp_path(self):
        # 清理文件:删除临时目录
        n = 1
        while True:
            try:
                rmtree(self.sys_tmp)
                return True
                break
            except Exception as e:
                msg = '清理临时目录({})操作失败：{}'.format(self.sys_tmp,e)
                sleep(1)
                if n > 10: break  # 偿试 n 次
                logging.error(msg)
                return False
            n += 1

    def run(self):
        logging.info('导入文件列表:'+  ('|').join(self.filelist))
        fn = len(self.filelist)
        for index,f in enumerate(self.filelist):
            msg = '【进程】{}/{} — {}'.format(index+1,fn,f)
            logging.info(msg)
            print(msg)

            if not self.updata_word_file(f):        # 更换当前导入文件
                msg = '读取文件失败：{}'.format(self.docx_file)
                print(msg)
                logging.warning(msg)
                self.run_info_register(sucess=False)
                continue

            try:
                self.update_dwdm()              # 更新单位代码库
                self.extract_from_word()        # 提取原始数据
                self.extract_from_word_clean()  # 清洗数据
                self.data2db_create()           # 提取精准数据（可以插入表中）
            except Exception as e:
                msg_warning = '提取数据失败：{}'.format(self.docx_file)
                print(msg_warning)
                logging.warning(msg_warning)
                self.run_info_register(sucess=False)
                continue
            else:
                pass

            run_result = self.import_data_to_table()     # 写入数据库
            if run_result != True:
                msg_warning = '写入数据库失败：{}'.format(self.docx_file)
                print(msg_warning)
                logging.warning(msg_warning)
                self.run_info_register(sucess=False)
                continue

            self.run_info_register(sucess=True)
            logging.info(msg + ' — 完成')

        self.update_db_dwdms()  # 将单位代码写入单位代码库

        # 控制中心最后的工作：将运行情况写入日志
        self.run_info_register(write_to_logfile = True)


class get_file_list():
    def __init__(self, path = '', ext=['.docx','.doc']):
        self.path = path
        self.ext = ext
        self.file_list = []
        self.get_file(self.path)

    def get_file(self, path):
        if os.path.isfile(path):    # 如果是文件
            if os.path.splitext(path)[1] in self.ext:
                self.file_list.append(path)
            return True

        list = os.listdir(path)
        for i in range(0,len(list)):
            file_or_dir = os.path.join(path,list[i])
            if os.path.isfile(file_or_dir):
                if os.path.splitext(file_or_dir)[1] in self.ext and '~$' not in file_or_dir:
                    self.file_list.append(file_or_dir)
            else:
                self.get_file(file_or_dir)

class get_dir_list():
    def __init__(self, path = ''):
        self.path = path
        self.dir_list = []
        self.get_dir(self.path)

    def get_dir(self, path):
        if os.path.isdir(path):    # 如果是文件夹
            self.dir_list.append(path)
            list = os.listdir(path)
            for i in list:
                self.get_dir(os.path.join(path,i))

if __name__ == '__main__':

    path = 'v:\\User\\人员档案资料'
    # path = r'v:\User\人员档案资料\02.直属单位'
    # path = r'v:\User\人员档案资料\02.直属单位\02.市政中心（15人）'
    # path = r'c:\temp'
    # path = r'v:\User\人员档案资料\01.机关\01.办公室（11人）'
    # path = r'v:\User\人员档案资料\02.直属单位\02.市政中心（15人）\徐珍珍.doc'
    path = r'c:\TEMP\tmp2'
    f = get_file_list(path=path)
    lt0 = f.file_list

    lt1 = ['emp_sample.docx','emp_xxb.docx']
    lt2 = ['emp_sample_word2003.doc','emp_sample_word2003_gif.doc','emp_sample_word2003_png.doc']
    lt5 = ['pb.doc']


    w = pickup_emp(lt0)

    import sys
    sys.exit()
